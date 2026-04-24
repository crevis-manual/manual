from __future__ import annotations

import hashlib
import html
import re
import shutil
from pathlib import Path
from typing import Iterable

import fitz
from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from PIL import Image


ROOT = Path(r"C:\Users\전영훈\Documents\manual")
TMP_ROOT = ROOT / ".tmp-scu"
ASCII_ROOT = Path(r"C:\manualtmp")
SCG_INDEX = ROOT / "manuals" / "camera" / "gige" / "SCG" / "index.html"
SCG_LOGO = ROOT / "manuals" / "camera" / "gige" / "SCG" / "images" / "crevis_logo.jpg"
DOCX_PATH = ASCII_ROOT / "SCU.docx"
PDF_PATH = ASCII_ROOT / "SCU.pdf"
OUT_DIR = ROOT / "manuals" / "camera" / "gige" / "SCU"
OUT_IMAGES = OUT_DIR / "images"
OUT_DOCX_IMAGES = OUT_IMAGES / "docx-extracted"
OUT_HTML = OUT_DIR / "index.html"
MANUAL_CENTER = ROOT / "index.html"


TOP_HEADING_RE = re.compile(r"^(?P<num>\d+)\.\s+(?P<title>\S.*)$")
SUB_HEADING_RE = re.compile(r"^(?P<num>\d+(?:\.\d+)+)\s+(?P<title>\S.*)$")
APPENDIX_RE = re.compile(r"^#\s*appendix\.[a-z]$", re.I)
URL_RE = re.compile(r"(https?://[^\s<]+)")


def normalize_key(text: str) -> str:
    lowered = clean_text(text).replace("·", "").replace("…", "").lower()
    return re.sub(r"[^0-9a-z가-힣]+", "", lowered)


def slugify(text: str, used: set[str]) -> str:
    slug = text.strip().lower()
    slug = re.sub(r"[^\w가-힣]+", "-", slug, flags=re.UNICODE)
    slug = re.sub(r"-{2,}", "-", slug).strip("-")
    if not slug:
        slug = "section"
    base = slug
    counter = 2
    while slug in used:
        slug = f"{base}-{counter}"
        counter += 1
    used.add(slug)
    return slug


def escape_and_link(text: str) -> str:
    def repl(match: re.Match[str]) -> str:
        url = match.group(1)
        escaped = html.escape(url)
        return f'<a href="{escaped}" target="_blank" rel="noopener">{escaped}</a>'

    return URL_RE.sub(repl, html.escape(text))


def clean_text(text: str) -> str:
    text = text.replace("\xa0", " ")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def split_toc_label(text: str) -> str:
    if "\t" in text:
        return text.split("\t", 1)[0].strip()
    return re.sub(r"[\.…·\s]+\d+\s*$", "", text).strip()


def extract_style_block() -> str:
    scg_html = SCG_INDEX.read_text(encoding="utf-8")
    match = re.search(r"<style>(.*?)</style>", scg_html, re.S)
    if not match:
        raise RuntimeError("Could not extract SCG style block")
    return match.group(1).strip()


def iter_block_items(parent: DocxDocument) -> Iterable[Paragraph | Table]:
    parent_elm = parent.element.body
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def part_extension(part) -> str:
    suffix = Path(str(part.partname)).suffix.lower()
    return suffix or ".png"


class ImageExtractor:
    def __init__(self, doc: Document, out_dir: Path) -> None:
        self.doc = doc
        self.out_dir = out_dir
        self.cache: dict[str, dict[str, str | int]] = {}
        self.out_dir.mkdir(parents=True, exist_ok=True)

    def extract_rid(self, rid: str) -> dict[str, str | int]:
        part = self.doc.part.related_parts[rid]
        digest = hashlib.sha1(part.blob).hexdigest()[:16]
        if digest in self.cache:
            return self.cache[digest]

        ext = part_extension(part)
        filename = f"{digest}{ext}"
        out_path = self.out_dir / filename
        if not out_path.exists():
            out_path.write_bytes(part.blob)

        width = height = 0
        try:
            with Image.open(out_path) as img:
                width, height = img.size
        except Exception:
            pass

        info = {
            "filename": filename,
            "path": f"images/docx-extracted/{filename}",
            "width": width,
            "height": height,
        }
        self.cache[digest] = info
        return info

    def paragraph_segments(self, paragraph: Paragraph) -> list[tuple[str, str | dict[str, str | int]]]:
        segments: list[tuple[str, str | dict[str, str | int]]] = []
        for run in paragraph.runs:
            if run.text:
                segments.append(("text", run.text))
            for blip in run._element.xpath(".//a:blip"):
                rid = blip.get(qn("r:embed"))
                if rid:
                    segments.append(("image", self.extract_rid(rid)))
        if not segments:
            for rid in paragraph._element.xpath(".//a:blip/@r:embed"):
                segments.append(("image", self.extract_rid(rid)))
        return segments


def render_inline_paragraph(text: str, segments: list[tuple[str, str | dict[str, str | int]]]) -> str:
    parts: list[str] = []
    for kind, value in segments:
        if kind == "text":
            if value:
                parts.append(escape_and_link(str(value)))
        else:
            info = value  # type: ignore[assignment]
            width = int(info["width"] or 0)
            height = int(info["height"] or 0)
            classes = ["inline-ui-img"]
            if width <= 60 or height <= 60:
                classes.append("icon")
            elif width >= 120 and height <= 80:
                classes.append("tabs")
            parts.append(f'<img src="{info["path"]}" alt="" class="{" ".join(classes)}">')
    html_text = "".join(parts).strip()
    if clean_text(text).startswith("- "):
        return f'<div class="inline-ui-line">{html_text}</div>'
    return f"<p>{html_text}</p>"


def render_table_cell_paragraph(paragraph: Paragraph, extractor: ImageExtractor) -> str:
    raw_text = paragraph.text or ""
    text = clean_text(raw_text)
    segments = extractor.paragraph_segments(paragraph)
    images = [value for kind, value in segments if kind == "image"]
    text_only = clean_text("".join(str(value) for kind, value in segments if kind == "text"))

    if images and not text_only:
        figures = "".join(
            f'<div class="table-cell-figure"><img src="{info["path"]}" alt="" class="table-cell-img"></div>'
            for info in images
        )
        return figures

    if images:
        if len(images) == 1:
            return (
                f'<div class="table-cell-figure"><img src="{images[0]["path"]}" alt="" class="table-cell-img"></div>'
                f'<p class="table-cell-caption">{escape_and_link(text_only)}</p>'
            )
        return render_inline_paragraph(raw_text, segments)

    if not text:
        return ""

    return f"<p>{escape_and_link(text)}</p>"


def render_table_cell(cell, extractor: ImageExtractor) -> tuple[str, bool]:
    parts: list[str] = []
    has_image = False
    for paragraph in cell.paragraphs:
        html_fragment = render_table_cell_paragraph(paragraph, extractor)
        if not html_fragment:
            continue
        if 'class="table-cell-figure"' in html_fragment or 'class="inline-ui-img' in html_fragment:
            has_image = True
        parts.append(html_fragment)
    if not parts:
        return "&nbsp;", has_image
    return f'<div class="table-cell-stack">{"".join(parts)}</div>', has_image


def render_table(table: Table, extractor: ImageExtractor) -> str:
    total_cells = 0
    image_cells = 0
    for row in table.rows:
        for cell in row.cells:
            total_cells += 1
            if any(paragraph._element.xpath(".//a:blip/@r:embed") for paragraph in cell.paragraphs):
                image_cells += 1

    has_images = image_cells > 0
    figure_table = has_images and image_cells * 2 >= max(total_cells, 1)

    rows: list[str] = []
    for row_index, row in enumerate(table.rows):
        tag = "th" if row_index == 0 and not has_images else "td"
        cells = []
        for cell in row.cells:
            cell_html, _ = render_table_cell(cell, extractor)
            cells.append(f"<{tag}>{cell_html}</{tag}>")
        rows.append(f"<tr>{''.join(cells)}</tr>")

    table_classes = ["docx-table"]
    if has_images:
        table_classes.append("has-media")
    if figure_table:
        table_classes.append("figure-table")

    return f'<div class="table-container"><table class="{" ".join(table_classes)}">{"".join(rows)}</table></div>'


def build_header_visual(doc: Document, extractor: ImageExtractor) -> str | None:
    candidates = [137, 124, 158, 160]
    for index in candidates:
        if index >= len(doc.paragraphs):
            continue
        segments = extractor.paragraph_segments(doc.paragraphs[index])
        images = [value for kind, value in segments if kind == "image"]
        if images:
            return str(images[0]["path"])
    return None


def parse_contact(doc: Document) -> dict[str, list[str] | str]:
    lines = [clean_text(p.text) for p in doc.paragraphs]
    start = lines.index("Contact")
    end = lines.index("Table of Contents")
    contact_lines = [line for line in lines[start:end] if line]

    sales: list[str] = []
    support: list[str] = []
    address = ""
    website = ""
    mode = None
    for line in contact_lines:
        if line == "Contact":
            continue
        if line.startswith("Product Inquiry"):
            mode = "sales"
            continue
        if line.startswith("Technical Support"):
            mode = "support"
            continue
        if line.startswith("29-4,"):
            address = line
            continue
        if line.startswith("www."):
            website = line
            continue
        if "=" in line:
            continue
        if mode == "sales":
            sales.append(line)
        elif mode == "support":
            support.append(line)

    return {
        "sales": sales,
        "support": support,
        "address": address,
        "website": website,
    }


def parse_cover_metadata() -> tuple[str, str, list[str]]:
    pdf = fitz.open(PDF_PATH)
    page1 = pdf[0].get_text("text")
    version_match = re.search(r"Version\s*:\s*([0-9.]+)", page1)
    release_match = re.search(r"Release Date\s*:\s*([0-9. ]+)", page1)
    version = version_match.group(1).strip() if version_match else "1.1"
    release = release_match.group(1).strip() if release_match else "2021. 03. 17"
    models = [line.strip(" -") for line in page1.splitlines() if line.strip().startswith("- ")]
    return version, release, models


def parse_toc(doc: Document) -> list[dict[str, str | int]]:
    entries = []
    for paragraph in doc.paragraphs:
        style = paragraph.style.name
        if style not in {"toc 1", "toc 2"}:
            continue
        label = split_toc_label(clean_text(paragraph.text))
        if not label:
            continue
        depth = 1 if style == "toc 1" else 2
        entries.append({"label": label, "depth": depth})
    return entries


def heading_depth(text: str) -> int | None:
    text = clean_text(text)
    if APPENDIX_RE.match(text) or text.lower() == "revision history":
        return 1
    top_match = TOP_HEADING_RE.match(text)
    if top_match:
        return 1
    sub_match = SUB_HEADING_RE.match(text)
    if sub_match:
        return sub_match.group("num").count(".") + 1
    return None


def render_main_content(doc: Document, extractor: ImageExtractor) -> tuple[list[dict[str, str]], dict[str, str]]:
    sections: list[dict[str, str]] = []
    heading_id_map: dict[str, str] = {}
    used_ids: set[str] = set()

    current_section: dict[str, str] | None = None
    current_parts: list[str] = []
    pending_plain: list[str] = []
    pending_images: list[dict[str, str | int]] = []
    started = False

    def flush_plain() -> None:
        nonlocal pending_plain
        if not pending_plain:
            return
        current_parts.append(f"<p>{escape_and_link(' '.join(pending_plain))}</p>")
        pending_plain = []

    def flush_images() -> None:
        nonlocal pending_images
        if not pending_images:
            return
        figures = []
        for info in pending_images:
            figures.append(
                f'<figure class="image-card"><img src="{info["path"]}" alt="" class="img-full"></figure>'
            )
        current_parts.append(f'<div class="image-grid">{"".join(figures)}</div>')
        pending_images = []

    def start_section(label: str) -> None:
        nonlocal current_section, current_parts
        flush_plain()
        flush_images()
        if current_section is not None:
            current_section["content"] = "".join(current_parts)
            sections.append(current_section)
        heading_id = slugify(label, used_ids)
        heading_id_map[normalize_key(label)] = heading_id
        current_section = {"id": heading_id, "title": label, "content": ""}
        current_parts = []

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            raw_text = block.text or ""
            text = clean_text(raw_text)
            depth = heading_depth(text) if text else None
            segments = extractor.paragraph_segments(block)
            images = [value for kind, value in segments if kind == "image"]
            text_only = clean_text("".join(str(value) for kind, value in segments if kind == "text"))
            image_only = bool(images) and not text_only

            if not started:
                if text == "1. 준비하기":
                    started = True
                    start_section(text)
                continue

            if depth == 1:
                if (
                    current_section is not None
                    and current_section["title"].lower().startswith("# appendix")
                    and not APPENDIX_RE.match(text)
                    and text.lower() != "revision history"
                ):
                    flush_plain()
                    flush_images()
                    heading_id = slugify(text, used_ids)
                    heading_id_map[normalize_key(text)] = heading_id
                    current_parts.append(f'<h3 id="{heading_id}">{html.escape(text)}</h3>')
                    if images:
                        pending_images.extend(images)
                    continue
                start_section(text)
                if images:
                    pending_images.extend(images)
                continue

            if current_section is None:
                continue

            if depth and depth >= 2:
                flush_plain()
                flush_images()
                heading_id = slugify(text, used_ids)
                heading_id_map[normalize_key(text)] = heading_id
                current_parts.append(f'<h3 id="{heading_id}">{html.escape(text)}</h3>')
                if images:
                    pending_images.extend(images)
                continue

            if image_only:
                flush_plain()
                pending_images.extend(images)
                continue

            if images:
                flush_plain()
                flush_images()
                current_parts.append(render_inline_paragraph(raw_text, segments))
                continue

            if not text:
                flush_plain()
                flush_images()
                continue

            if text.startswith("※"):
                flush_plain()
                flush_images()
                current_parts.append(f'<p class="inline-note-title warning">{escape_and_link(text)}</p>')
                continue

            if text.startswith("▶"):
                flush_plain()
                flush_images()
                item = clean_text(text[1:])
                current_parts.append(f'<ul class="feature-list"><li>{escape_and_link(item)}</li></ul>')
                continue

            if text.startswith("- "):
                flush_plain()
                flush_images()
                current_parts.append(f"<p><strong>{escape_and_link(text)}</strong></p>")
                continue

            if re.match(r"^\d+\)", text):
                flush_plain()
                flush_images()
                current_parts.append(f"<p>{escape_and_link(text)}</p>")
                continue

            pending_plain.append(text)

        else:
            if not started or current_section is None:
                continue
            flush_plain()
            flush_images()
            current_parts.append(render_table(block, extractor))

    flush_plain()
    flush_images()
    if current_section is not None:
        current_section["content"] = "".join(current_parts)
        sections.append(current_section)

    return sections, heading_id_map


def build_sidebar(toc_entries: list[dict[str, str | int]], heading_id_map: dict[str, str]) -> str:
    toc_aliases = {
        normalize_key("3.3 USB3 Connector"): normalize_key("3.3 USB3.1 Gen.1 Micro-B Type Connector"),
        normalize_key("4.9 Chunk Control"): normalize_key("4.9 Chunk Data Control"),
        normalize_key("4.15 Defect Pixel Control"): normalize_key("4.15 Defect Pixel Correction"),
    }

    items: list[str] = []
    for entry in toc_entries:
        label = str(entry["label"])
        depth = int(entry["depth"])
        key = normalize_key(label)
        target = heading_id_map.get(key)
        if not target and key in toc_aliases:
            target = heading_id_map.get(toc_aliases[key])
        if not target:
            continue
        css_class = ' class="toc-sub"' if depth > 1 else ""
        items.append(f'<li{css_class}><a href="#{target}"><span>{html.escape(label)}</span></a></li>')
    return "\n".join(items)


def render_contact_section(contact: dict[str, list[str] | str]) -> str:
    sales_lines = "".join(f"<p>{escape_and_link(line)}</p>" for line in contact["sales"])  # type: ignore[index]
    support_lines = "".join(f"<p>{escape_and_link(line)}</p>" for line in contact["support"])  # type: ignore[index]
    address = escape_and_link(str(contact["address"]))
    website = escape_and_link(str(contact["website"]))
    return f"""
<section id="contact">
    <h2>Contact</h2>
    <div class="summary-grid">
        <article class="summary-card">
            <strong>Product Inquiry (Sales)</strong>
            {sales_lines}
        </article>
        <article class="summary-card">
            <strong>Technical Support</strong>
            {support_lines}
        </article>
        <article class="summary-card">
            <strong>Address</strong>
            <p>{address}</p>
            <p>{website}</p>
        </article>
    </div>
</section>
""".strip()


def render_sections(sections: list[dict[str, str]]) -> str:
    rendered = []
    for section in sections:
        rendered.append(
            f'<section id="{section["id"]}"><h2>{html.escape(section["title"])}</h2>{section["content"]}</section>'
        )
    return "\n".join(rendered)


def manual_center_card_exists(index_html: str) -> bool:
    return "manuals/camera/gige/SCU/index.html" in index_html


def ensure_manual_center_entry() -> None:
    index_html = MANUAL_CENTER.read_text(encoding="utf-8")
    if manual_center_card_exists(index_html):
        return

    insertion = """
        <a href="manuals/camera/gige/SCU/index.html" class="manual-card" data-category="USB" data-name="SCU제품매뉴얼 SCU Manual USB3Vision Camera USER MANUAL Sony CMOS Sensor" target="_blank" rel="noopener">
            <h3>SCU 제품매뉴얼</h3>
            <p>USB3Vision Camera USER MANUAL</p>
            <span class="tag usb">USB 3.0</span>
        </a>
""".rstrip()

    marker = """        <a href="manuals/camera/gige/SCG/index_en.html" class="manual-card" data-category="GigE" data-name="SCG제품매뉴얼 SCG Manual Korean English Sony CMOS GigE Camera USER MANUAL" target="_blank" rel="noopener">"""
    if marker not in index_html:
        raise RuntimeError("Could not find SCG manual card in manual center")

    index_html = index_html.replace(marker, insertion + "\n\n" + marker, 1)
    MANUAL_CENTER.write_text(index_html, encoding="utf-8")


def main() -> None:
    OUT_DOCX_IMAGES.mkdir(parents=True, exist_ok=True)
    OUT_IMAGES.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SCG_LOGO, OUT_IMAGES / "crevis_logo.jpg")

    style_block = extract_style_block()
    doc = Document(DOCX_PATH)
    extractor = ImageExtractor(doc, OUT_DOCX_IMAGES)

    version, release_date, models = parse_cover_metadata()
    toc_entries = parse_toc(doc)
    contact = parse_contact(doc)
    sections, heading_id_map = render_main_content(doc, extractor)
    sidebar_html = build_sidebar(toc_entries, heading_id_map)
    contact_html = render_contact_section(contact)
    sections_html = render_sections(sections)
    header_visual = build_header_visual(doc, extractor)

    model_items = "".join(f"<li>{html.escape(model)}</li>" for model in models)
    header_visual_html = (
        f'<div class="header-visual-group"><img src="images/crevis_logo.jpg" class="header-logo" alt="Crevis Logo"><img src="{header_visual}" class="header-camera-main scu-header-visual" alt="SCU visual"></div>'
        if header_visual
        else '<div class="header-visual-group"><img src="images/crevis_logo.jpg" class="header-logo" alt="Crevis Logo"></div>'
    )

    extra_css = """
        .hero-models {
            margin: 18px 0 0;
            padding-left: 18px;
            display: grid;
            gap: 8px;
            color: #2b4250;
        }
        .hero-models li {
            line-height: 1.45;
        }
        .summary-card p {
            margin: 0 0 8px;
            line-height: 1.6;
        }
        .summary-card p:last-child {
            margin-bottom: 0;
        }
        .scu-header-visual {
            max-width: 240px;
            background: #fff;
            border-radius: 18px;
            padding: 14px;
            border: 1px solid #dde9ec;
            box-shadow: 0 10px 24px rgba(22, 57, 74, 0.06);
        }
        .content > section > .table-container:first-child,
        .content > section > .image-grid:first-child {
            margin-top: 4px;
        }
        h3 {
            gap: 0;
            padding-left: 14px;
        }
        h3::before {
            content: none !important;
            display: none !important;
        }
        .docx-table.has-media td,
        .docx-table.has-media th {
            vertical-align: top;
        }
        .docx-table.figure-table td,
        .docx-table.figure-table th {
            text-align: center;
        }
        .table-cell-stack {
            display: grid;
            gap: 10px;
        }
        .docx-table.figure-table .table-cell-stack {
            justify-items: center;
            text-align: center;
        }
        .table-cell-stack > p {
            margin: 0;
            line-height: 1.6;
        }
        .table-cell-figure {
            display: flex;
            justify-content: center;
            align-items: center;
        }
        .table-cell-img {
            display: block;
            max-width: 100%;
            width: auto;
            height: auto;
            max-height: 260px;
            object-fit: contain;
            margin: 0 auto;
        }
        .table-cell-caption {
            margin: 0;
            line-height: 1.55;
        }
        .docx-table.figure-table .table-cell-caption {
            text-align: center;
        }
    """

    html_output = f"""<!DOCTYPE html>
<html lang="ko">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>SCU 제품매뉴얼 | Crevis</title>
    <link rel="stylesheet" href="../../../../css/product_manual/style.css">
    <style>
{style_block}
{extra_css}
    </style>
</head>
<body>
<nav class="sidebar">
    <p class="eyebrow">CREVIS MANUAL CENTER</p>
    <h2>SCU Manual</h2>
    <a class="back-link" href="../../../../index.html"><span>&larr; Manual Center</span></a>
    <p class="rev">Version {html.escape(version)}</p>
    <ul class="toc">
{sidebar_html}
    </ul>
</nav>

<div class="content">
    <header class="header-box">
        <div class="header-text">
            <p class="category-label">USB3Vision Camera USER MANUAL</p>
            <h1 class="model-title">SCU 제품매뉴얼</h1>
            <p style="font-size: 1.2rem; color: #6b7d87; margin-top: 10px;">Sony CMOS Sensor USB3Vision Camera USER MANUAL</p>
            <div class="meta-strip">
                <span class="pill">Version {html.escape(version)}</span>
                <span class="pill">Release Date: {html.escape(release_date)}</span>
                <span class="pill">Source: PDF + DOCX conversion</span>
            </div>
            <ul class="hero-models">
                {model_items}
            </ul>
        </div>
        {header_visual_html}
    </header>
    {contact_html}
    {sections_html}
</div>
</body>
</html>
"""

    OUT_HTML.write_text(html_output, encoding="utf-8")
    ensure_manual_center_entry()


if __name__ == "__main__":
    main()
